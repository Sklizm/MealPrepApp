"""
Continutul (textul) raportului de practica, in limba romana.

Fiecare functie primeste documentul `doc` si modulul de ajutoare `h`
(docx_helpers) si adauga o sectiune. Listingurile de cod din Anexe sunt citite
direct din fisierele reale ale proiectului la momentul generarii, deci raman
mereu sincronizate cu codul.
"""

import os

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def read_repo_file(relpath, max_lines=None):
    path = os.path.join(REPO, relpath)
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    if max_lines:
        lines = text.split("\n")
        if len(lines) > max_lines:
            text = "\n".join(lines[:max_lines]) + "\n[... continua in fisierul sursa ...]"
    return text


# =========================================================================== #
#  1. FOAIE DE TITLU  (placeholder — se inlocuieste cu modelul institutiei)
# =========================================================================== #

def title_page(doc, h):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    def centered(text, size=12, bold=False, space_before=0, space_after=0, caps=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text.upper() if caps else text)
        r.font.name = h.BODY_FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        return p

    centered("[MINISTERUL EDUCATIEI SI CERCETARII]", 12, bold=True, space_before=0)
    centered("[DENUMIREA INSTITUTIEI DE INVATAMANT]", 14, bold=True, space_after=2)
    centered("[Specialitatea / Programul de studii]", 12, space_after=120)

    centered("RAPORT", 16, bold=True, space_after=2)
    centered("privind stagiul de practica", 14, bold=True, space_after=4)
    centered("Tema: Aplicatie de gestionare a retetelor si planificare a meselor "
             "(„MealPrep”) — baza de date SQL Server si aplicatie desktop WPF",
             13, bold=True, space_after=120)

    # Bloc dreapta: student + coordonator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.5
    for line, bold in [
        ("A elaborat: [Nume Prenume student]", False),
        ("Grupa: [grupa]", False),
        ("Conducator de practica: [Nume Prenume]", False),
        ("Locul desfasurarii practicii: [intreprindere / catedra]", False),
    ]:
        r = p.add_run(line + "\n")
        r.font.name = h.BODY_FONT
        r.font.size = Pt(12)
        r.font.bold = bold

    centered("[Oras]   [An]", 12, bold=True, space_before=160)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run("(Foaie de titlu provizorie — se inlocuieste cu modelul oficial al institutiei.)")
    rn.font.name = h.BODY_FONT
    rn.font.size = Pt(9)
    rn.font.italic = True


# =========================================================================== #
#  2. CUPRINS
# =========================================================================== #

def cuprins(doc, h):
    h.heading(doc, "Cuprins", level=1)
    h.toc_field(doc)


# =========================================================================== #
#  3. INTRODUCERE
# =========================================================================== #

def introducere(doc, h):
    h.heading(doc, "Introducere", level=1)
    h.body(doc,
        "Prezentul raport descrie activitatea desfasurata in cadrul stagiului de practica, "
        "constand in proiectarea si dezvoltarea aplicatiei „MealPrep” — un sistem informatic "
        "destinat gestionarii retetelor culinare si planificarii meselor. Lucrarea acopera ambele "
        "componente ale produsului: o baza de date relationala SQL Server, gazduita intr-un container "
        "Docker, si o aplicatie desktop pentru Windows realizata in tehnologia WPF (.NET 10), cu interfata "
        "in limba romana.")
    h.body(doc,
        "Tema a fost aleasa pentru ca imbina, intr-un produs realist si demonstrabil, cele mai importante "
        "competente vizate de stagiu: modelarea datelor, securitatea accesului la baza de date, "
        "arhitectura unei aplicatii cu interfata grafica si buna practica de versionare si documentare a "
        "codului. Aplicatia raspunde unei nevoi concrete — organizarea retetelor, a stocului din camara si "
        "a planului de mese saptamanal, cu generarea automata a listei de cumparaturi.")

    h.heading(doc, "Obiectivele lucrarii", level=2)
    h.body(doc, "Obiectivele propuse la inceputul stagiului au fost urmatoarele:")
    h.numbered(doc, "O1. Proiectarea unei baze de date relationale normalizate, cu integritate referentiala "
                    "asigurata prin chei externe si constrangeri explicite.")
    h.numbered(doc, "O2. Expunerea unui API de date sigur, exclusiv prin proceduri stocate, pe baza "
                    "principiului privilegiului minim (least privilege), astfel incat aplicatia sa nu poata "
                    "accesa direct tabelele.")
    h.numbered(doc, "O3. Dezvoltarea unei aplicatii desktop WPF, structurata dupa sablonul arhitectural "
                    "MVVM, care comunica cu baza de date prin acest API.")
    h.numbered(doc, "O4. Implementarea functionalitatilor de baza: autentificare cu blocare la atacuri de "
                    "tip brute-force, gestionarea retetelor, a ingredientelor si a camarii, lista de "
                    "cumparaturi calculata, planificarea meselor si rapoarte.")
    h.numbered(doc, "O5. Aplicarea unui flux de lucru profesional de documentare (Obsidian) si versionare "
                    "(Git), folosind un asistent de programare bazat pe inteligenta artificiala (Claude Code).")
    h.body(doc,
        "Structura raportului urmeaza cerintele indrumarului: dupa aceasta introducere sunt descrise modul "
        "de elaborare a produsului program, listingul codului sursa si rezultatele testarii (date de intrare, "
        "date de iesire si functionalitate), urmate de observatii generale si concluzie, bibliografie si anexe. "
        "Anexele includ si ghiduri reproductibile privind mediul de lucru (asistentul Claude Code, Obsidian si Git).")


# =========================================================================== #
#  4. CONTINUTUL ACTIVITATILOR SI SARCINILOR DE LUCRU
# =========================================================================== #

def continut(doc, h):
    h.heading(doc, "Continutul activitatilor si sarcinilor de lucru", level=1)

    # ---- 4.a Descrierea modului de elaborare ---- #
    h.heading(doc, "Descrierea modului de elaborare a produsului program", level=2)

    h.heading(doc, "Arhitectura generala", level=3)
    h.body(doc,
        "Produsul este format din doua jumatati care comunica printr-un contract strict. Baza de date "
        "SQL Server expune un set de proceduri stocate; aplicatia WPF le apeleaza, autentificandu-se cu un "
        "cont SQL cu privilegii minime (mealprep_app). Aplicatia nu are si nu poate avea acces direct la "
        "tabele — orice citire sau modificare trece printr-o procedura. Aceasta separare aduce doua avantaje: "
        "imposibilitatea structurala a injectiei SQL din partea aplicatiei si un contract clar, versionabil, "
        "intre baza de date si interfata.")
    h.figure_placeholder(doc, "Arhitectura pe straturi: WPF (View/ViewModel) → repository-uri Dapper → "
                               "proceduri stocate → tabele SQL Server (in Docker).")

    h.heading(doc, "Tehnologii utilizate", level=3)
    h.body(doc, "Componenta de baza de date:")
    h.bullet(doc, "SQL Server 2022, rulat intr-un container Docker (port 1433), construit prin scripturi T-SQL idempotente.")
    h.bullet(doc, "sqlcmd (mssql-tools18) pentru rularea scriptata a build-ului; DataGrip ca instrument GUI.")
    h.body(doc, "Componenta de aplicatie:")
    h.bullet(doc, "WPF pe .NET 10 (net10.0-windows), sablon arhitectural MVVM.")
    h.bullet(doc, "CommunityToolkit.Mvvm pentru ObservableProperty si RelayCommand (reducerea codului repetitiv).")
    h.bullet(doc, "Dapper peste Microsoft.Data.SqlClient pentru accesul la date (apel de proceduri stocate).")
    h.bullet(doc, "Microsoft.Extensions.DependencyInjection pentru injectia de dependinte.")
    h.bullet(doc, "BCrypt.Net pentru parolarea (hashing) parolelor; ClosedXML pentru exportul in Excel.")

    h.heading(doc, "Proiectarea bazei de date", level=3)
    h.body(doc,
        "Baza de date MealPrepDB contine 12 tabele: sase de baza (Users, Units, Categories, Ingredients, "
        "Recipes, RecipeIngredients), doua pentru securitate si audit (PasswordHistory, AuditLog), trei "
        "pentru planificarea meselor (MealPlanEntries, RecipeFavorites, UserPantry) si una de tip lookup "
        "(IngredientCategories). Ordinea de construire respecta dependentele dintre tabele si este codificata "
        "in scriptul master run_all.sql.")
    h.body(doc, "Conventiile de proiectare aplicate consecvent (motivate in jurnalul de decizii):")
    h.bullet(doc, "Scripturi idempotente: fiecare CREATE este protejat (IF OBJECT_ID(...) IS NULL), iar "
                  "seed-urile folosesc MERGE, astfel incat re-rularea build-ului nu distruge datele locale.")
    h.bullet(doc, "Toate sirurile de caractere sunt NVARCHAR (Unicode); marcajele temporale sunt UTC, prin "
                  "valoarea implicita SYSUTCDATETIME().")
    h.bullet(doc, "Numele constrangerilor sunt explicite, cu prefixe PK_, FK_, UQ_, CK_, DF_, IX_ — lizibile "
                  "in mesajele de eroare si stabile la reconstructie.")
    h.bullet(doc, "Stergerile in cascada sunt minime, intentionat: singura cascada principala este "
                  "Recipes → RecipeIngredients; restul relatiilor sunt RESTRICT, pentru a nu distruge date "
                  "din neatentie.")
    h.body(doc,
        "Securitatea accesului este realizata prin contul mealprep_app, membru al rolului mealprep_app_role, "
        "care are GRANT EXECUTE pe schema dbo, dar DENY pe SELECT/INSERT/UPDATE/DELETE. Modificarile reusesc "
        "doar prin proceduri, gratie mecanismului de ownership chaining (procedurile si tabelele au acelasi "
        "proprietar, dbo). Contul administrativ sa este rezervat exclusiv migrarilor.")
    h.body(doc, "Mecanisme de securitate si fiabilitate implementate la nivel de baza de date:")
    h.bullet(doc, "Blocare la autentificare: dupa 5 esecuri consecutive, contul este blocat 15 minute "
                  "(eveniment consemnat in AuditLog).")
    h.bullet(doc, "Istoric de parole: ultimele 5 valori (hash) sunt pastrate per utilizator; reutilizarea "
                  "este respinsa (eroare 50001).")
    h.bullet(doc, "Concurenta optimista pe Recipes: o coloana RowVersion este returnata la citire si "
                  "verificata la actualizare; o versiune invechita produce eroarea 50004.")
    h.bullet(doc, "Jurnal de audit: fiecare procedura care modifica date scrie o linie in AuditLog, in "
                  "aceeasi tranzactie.")
    h.body(doc,
        "Pentru transferul listelor de date intre aplicatie si proceduri s-au folosit doua mecanisme: JSON "
        "(parsat cu OPENJSON) pentru sarcinile de scriere (lista de ingrediente la crearea unei retete) si "
        "un parametru de tip tabel (TVP dbo.IntList) pentru filtrele de citire. Lista de cumparaturi nu este "
        "stocata, ci calculata la cerere de procedura sp_GetShoppingList, prin imbinarea meselor planificate "
        "cu ingredientele retetelor, scalate dupa numarul de portii, minus stocul existent in camara.")

    h.body(doc, "Dezvoltarea bazei de date s-a desfasurat pe patru faze:")
    h.bullet(doc, "Faza 1 — schema de baza (6 tabele) si seed-urile pentru unitati si categorii.")
    h.bullet(doc, "Faza 2 si 2.5 — stratul de securitate (login cu privilegii minime, audit, istoric parole, "
                  "blocare), indexarea cheilor externe si concurenta optimista; in total 18+ proceduri.")
    h.bullet(doc, "Faza 3 — planificarea meselor, favoritele, camara si lista de cumparaturi calculata.")
    h.bullet(doc, "Faza 4 — categorisirea ingredientelor (lookup), procedurile de raportare si o procedura "
                  "sigura de citire a profilului (fara expunerea hash-ului parolei).")

    h.heading(doc, "Dezvoltarea aplicatiei WPF", level=3)
    h.body(doc,
        "Aplicatia este organizata dupa sablonul MVVM: vederile (View, fisiere XAML) nu contin logica de "
        "business, ci se leaga prin data binding de ViewModel-uri. ViewModel-urile folosesc atributele "
        "din CommunityToolkit.Mvvm (ObservableProperty, RelayCommand) si apeleaza repository-uri specializate "
        "(RecipeRepository, IngredientRepository, MealPlanRepository etc.), fiecare incapsuland apelurile de "
        "proceduri prin Dapper. Serviciile transversale (navigare, dialoguri, sesiune, hashing de parole) "
        "sunt injectate prin containerul de dependinte configurat in App.xaml.cs.")
    h.body(doc,
        "Interfata foloseste un sistem de design unitar — paleta crem / masliniu / maro-inchis, definita in "
        "Themes/Colors.xaml si Themes/Styles.xaml. Ferestrele sunt fara chrome nativ Windows (clasa "
        "WindowChrome), cu bara de titlu proprie; casetele de mesaj native (MessageBox) au fost inlocuite cu "
        "un dialog stilizat unic (MessageDialog) cu trei variante: informare, confirmare si eroare. Controale "
        "native precum DatePicker, Calendar, Menu, ToolTip si ScrollBar au fost retematizate prin stiluri "
        "implicite (fara cheie), astfel incat intreaga aplicatie ramane consecventa vizual.")
    h.body(doc, "Aplicatia a fost construita incremental, pe ecranele descrise in specificatia de design:")
    h.bullet(doc, "Autentificare — inregistrare, conectare, profil, schimbarea parolei.")
    h.bullet(doc, "Acasa — tablou de bord cu indicatori (KPI) si retete recente.")
    h.bullet(doc, "Retete — lista, vederea de detaliu si editorul (ingrediente introduse ca tabel, "
                  "salvare protejata de concurenta optimista).")
    h.bullet(doc, "Ingrediente — lista (plata sau grupata pe categorii) cu cautare instantanee; Frigider "
                  "(camara) cu adaugare/editare/stergere; Lista de cumparaturi cu interval de date, export "
                  "Excel si tiparire.")
    h.bullet(doc, "Planificare — calendar de mese in doua vederi (lunar 6×7 si saptamanal 7×4), cu dialog "
                  "de adaugare/editare/stergere a unei mese si scurtatura „Adauga la plan” din detaliul retetei.")
    h.bullet(doc, "Rapoarte — trei sub-module: statistici lunare (indicatori, defalcare pe categorii de masa, "
                  "top retete si top ingrediente), plan saptamanal pentru tiparire si lista de cumparaturi pentru "
                  "tiparire, ambele cu tiparire si export Excel.")

    h.heading(doc, "Metodologia de lucru", level=3)
    h.body(doc,
        "Proiectul a fost realizat in mai multe sesiuni de lucru. Pentru continuitate s-a folosit un vault "
        "Obsidian ca sursa unica de adevar: jurnalul de decizii arhitecturale (append-only), note per tabel, "
        "specificatia de design si jurnalele de sesiune. Fiecare nota are un corespondent in limba romana "
        "(fisier -ro.md). Codul este versionat in Git, lucrand pe ramuri de functionalitate cu commit-uri si "
        "push-uri frecvente, mici si descriptive. Pe parcurs a fost folosit asistentul de programare Claude "
        "Code, care a accelerat scrierea codului, documentarea si diagnosticarea. Detalii reproductibile despre "
        "aceste instrumente se gasesc in Anexele A2–A4.")

    # ---- 4.b Listingul produsului program ---- #
    h.heading(doc, "Listingul produsului program", level=2)
    h.body(doc,
        "In continuare sunt prezentate fragmente reprezentative din codul sursa, cu rol ilustrativ. "
        "Listingul complet (toate scripturile de baza de date si fisierele esentiale ale aplicatiei) se "
        "gaseste in Anexa A1.")

    h.heading(doc, "Definirea tabelelor Recipes si RecipeIngredients (T-SQL)", level=3)
    h.body(doc, "Tabelul de jonctiune RecipeIngredients ilustreaza conventiile de proiectare: cheie unica "
                "(reteta, ingredient), constrangere CHECK pe cantitate, cascada catre reteta, RESTRICT pe "
                "unitate si ingredient.")
    h.code_block(doc, read_repo_file("Database/06_recipe_ingredients.sql"))

    h.heading(doc, "Procedura sp_CreateRecipe — scriere tranzactionala cu OPENJSON", level=3)
    h.body(doc, "Creeaza reteta si liniile de ingrediente intr-o singura tranzactie; lista de ingrediente "
                "soseste ca JSON si este parsata cu OPENJSON.")
    h.code_block(doc, _extract_proc(read_repo_file("Database/procs/02_recipes_write.sql"),
                                    "sp_CreateRecipe"))

    h.heading(doc, "Maparea erorilor de baza de date in mesaje pentru utilizator (C#)", level=3)
    h.body(doc, "Codurile de eroare SQL (proprii 50000–50004 si native) sunt traduse in mesaje prietenoase "
                "in limba romana, pentru ca nicio eroare sa nu ajunga opaca la utilizator.")
    h.code_block(doc, read_repo_file("App/MealPrepApp/Data/DbExceptionMapper.cs"))

    h.heading(doc, "ViewModel: salvarea unei retete cu validare si concurenta optimista (C#)", level=3)
    h.body(doc, "Metoda Save din ReteteEditorViewModel valideaza campurile, blocheaza ingredientele "
                "duplicate inainte de salvare si trateaza conflictul de editare (eroarea 50004).")
    h.code_block(doc, _extract_method_save(read_repo_file(
        "App/MealPrepApp/ViewModels/Retete/ReteteEditorViewModel.cs")))

    # ---- 4.c Rezultatele testarii ---- #
    h.heading(doc, "Rezultatele testarii produsului program", level=2)
    h.body(doc,
        "Testarea a vizat doua niveluri: (1) comportamentul bazei de date (constrangeri, securitate, "
        "proceduri), verificat prin sqlcmd direct in container, si (2) functionalitatea aplicatiei, verificata "
        "manual pe o masina virtuala Windows 11 (aplicatia WPF nu se poate compila pe statia Linux de "
        "dezvoltare). Mai jos sunt grupate datele de intrare, datele de iesire corespunzatoare si descrierea "
        "functionalitatii.")

    h.heading(doc, "Datele de intrare si datele de iesire", level=3)
    h.body(doc, "Tabelul urmator rezuma cazurile de test, cu intrarea aplicata si iesirea observata.")
    h.simple_table(
        doc,
        ["Caz de test", "Date de intrare", "Date de iesire (rezultat)"],
        [
            ["Cantitate invalida la ingredient",
             "Cantitate = 0 sau negativa",
             "Respins de CK_RecipeIngredients_Quantity; randul nu se salveaza."],
            ["Ingredient duplicat in reteta",
             "Acelasi ingredient pe doua randuri",
             "Blocat in editor cu mesaj clar; in absenta verificarii ar produce eroarea SQL 2627 (UNIQUE)."],
            ["Reteta cu 18 ingrediente distincte",
             "18 ingrediente diferite + descriere + instructiuni",
             "Salvare reusita; procedura returneaza noul RecipeID."],
            ["Stergerea unei retete",
             "DELETE pe o reteta cu linii de ingrediente",
             "Liniile din RecipeIngredients se sterg automat (ON DELETE CASCADE)."],
            ["Stergerea unui ingredient folosit",
             "DELETE pe un ingredient prezent in retete",
             "Blocat (RESTRICT) — integritatea referentiala este pastrata."],
            ["Blocare la autentificare",
             "5 incercari de conectare cu parola gresita",
             "Cont blocat 15 minute; eveniment ACCOUNT_LOCKED in AuditLog."],
            ["Reutilizarea parolei",
             "Parola noua = una din ultimele 5",
             "Respins cu eroarea 50001 (mesaj: parola folosita recent)."],
            ["Conflict de editare (concurenta)",
             "Salvare cu un RowVersion invechit",
             "Eroarea 50004; aplicatia ofera reincarcarea retetei."],
            ["Lista de cumparaturi",
             "Plan de mese pe un interval + stoc partial in camara",
             "Cantitati de cumparat scalate dupa portii, minus stocul existent."],
            ["Planificarea unei mese",
             "Reteta + data + slot (mic dejun/pranz/cina/gustare) din dialogul de plan",
             "Masa apare in calendarul lunar si saptamanal; se poate edita sau sterge."],
            ["Statistici lunare",
             "Selectarea unei luni cu mese planificate",
             "Total mese, retete/ingrediente distincte, top 5 retete si top 10 ingrediente."],
        ],
        caption="Tabelul 1 — Cazuri de test: date de intrare si date de iesire."
    )

    h.heading(doc, "Functionalitatea produsului program", level=3)
    h.body(doc,
        "Functionalitatea a fost verificata pe ecranele principale ale aplicatiei. Mai jos sunt rezervate "
        "spatii pentru capturile de ecran realizate pe masina virtuala Windows, fiecare cu legenda "
        "corespunzatoare.")
    h.figure_placeholder(doc, "Ecranul de autentificare (Conectare / Inregistrare).")
    h.figure_placeholder(doc, "Tabloul de bord (Acasa) cu indicatori si retete recente.")
    h.figure_placeholder(doc, "Lista de retete si vederea de detaliu a unei retete.")
    h.figure_placeholder(doc, "Editorul de reteta, cu tabelul editabil de ingrediente.")
    h.figure_placeholder(doc, "Modulul Ingrediente: Frigider (camara) si Lista de cumparaturi.")
    h.figure_placeholder(doc, "Modulul Planificare (calendar lunar / saptamanal) si dialogul de adaugare masa.")
    h.figure_placeholder(doc, "Modulul Rapoarte: statistici lunare (indicatori, top retete si ingrediente).")
    h.figure_placeholder(doc, "Modulul Rapoarte: plan saptamanal / lista de cumparaturi pentru tiparire.")
    h.body(doc,
        "In urma testarii, toate fluxurile aplicatiei (autentificare, gestionarea retetelor, a ingredientelor "
        "si a camarii, generarea listei de cumparaturi, planificarea meselor in calendar si rapoartele lunare) "
        "functioneaza conform cerintelor, confirmate pe masina virtuala Windows 11. Mesajele de eroare sunt "
        "afisate stilizat, in limba romana, fara blocarea aplicatiei.")


def _extract_proc(sql_text, proc_name):
    """Extrage blocul unei singure proceduri dintr-un fisier .sql (pana la GO-ul final al ei)."""
    lines = sql_text.split("\n")
    out, capturing = [], False
    for i, line in enumerate(lines):
        upper = line.upper()
        if (f"PROCEDURE DBO.{proc_name.upper()}" in upper) or (f"PROCEDURE {proc_name.upper()}" in upper):
            # urca pana la linia CREATE
            j = i
            while j > 0 and "CREATE" not in lines[j].upper():
                j -= 1
            out = lines[j:i]
            capturing = True
        if capturing:
            out.append(line)
            if line.strip().upper() == "GO" and len(out) > 5:
                break
    return "\n".join(out) if out else sql_text


def _extract_method_save(cs_text):
    """Extrage metoda Save() (cu atributul [RelayCommand]) din ReteteEditorViewModel."""
    lines = cs_text.split("\n")
    start = None
    for i, line in enumerate(lines):
        if "private async Task Save()" in line:
            start = i - 1 if "[RelayCommand]" in lines[i - 1] else i
            break
    if start is None:
        return cs_text
    depth, end = 0, None
    seen_brace = False
    for i in range(start, len(lines)):
        depth += lines[i].count("{") - lines[i].count("}")
        if "{" in lines[i]:
            seen_brace = True
        if seen_brace and depth == 0:
            end = i
            break
    return "\n".join(lines[start:(end + 1) if end else len(lines)])


# =========================================================================== #
#  5. OBSERVATII GENERALE. CONCLUZIE
# =========================================================================== #

def concluzie(doc, h):
    h.heading(doc, "Observatii generale. Concluzie", level=1)
    h.body(doc,
        "Stagiul de practica s-a finalizat cu un produs program functional, care acopera obiectivele propuse "
        "in introducere. A fost realizata o baza de date relationala normalizata, cu integritate referentiala "
        "asigurata prin constrangeri explicite (O1), expusa printr-un API de proceduri stocate pe principiul "
        "privilegiului minim, ceea ce face injectia SQL imposibila structural din partea aplicatiei (O2). "
        "Peste acest API a fost dezvoltata o aplicatie desktop WPF, organizata dupa sablonul MVVM (O3), care "
        "implementeaza autentificarea cu blocare anti-brute-force, gestionarea retetelor, a ingredientelor si "
        "a camarii, lista de cumparaturi calculata si modulele de planificare si raportare (O4). Intregul "
        "proces a fost documentat in Obsidian si versionat in Git, cu sprijinul unui asistent AI (O5).")
    h.body(doc,
        "Importanta lucrarii consta in faptul ca reproduce, la scara unui proiect didactic, deciziile reale "
        "dintr-o aplicatie de productie: separarea stricta intre date si interfata, securitatea bazata pe "
        "privilegii minime, concurenta optimista, auditarea actiunilor si un sistem de design coerent.")

    h.heading(doc, "Complexitate si dificultati intampinate", level=2)
    h.body(doc,
        "Cercetarea nu a fost lipsita de dificultati, iar acestea merita mentionate pentru ca arata limitele "
        "actuale ale lucrarii:")
    h.bullet(doc, "Aplicatia WPF (net10.0-windows) nu poate fi compilata pe statia de dezvoltare Linux; "
                  "verificarea functionala s-a facut pe o masina virtuala Windows 11, ceea ce a incetinit "
                  "ciclul de testare.")
    h.bullet(doc, "Camara nu realizeaza conversii intre unitati de masura (de exemplu grame ↔ cani); stocul "
                  "este urmarit exact pe tuplul (utilizator, ingredient, unitate).")
    h.bullet(doc, "O eroare la salvarea unei retete (un ingredient introdus de doua ori) ajungea la utilizator "
                  "ca mesaj generic. Cauza reala — incalcarea cheii unice (eroarea SQL 2627) — a fost greu de "
                  "identificat tocmai pentru ca mesajul era opac. Solutia a fost dubla: o verificare anti-duplicat "
                  "in editor, inainte de salvare, si maparea explicita a codurilor SQL in mesaje clare, cu afisarea "
                  "codului brut pentru erorile neanticipate.")

    h.heading(doc, "Directii de dezvoltare viitoare", level=2)
    h.body(doc, "Plecand chiar de la ceea ce nu a fost realizat in versiunea actuala, se contureaza urmatoarele "
                "directii de continuare:")
    h.bullet(doc, "Adaugarea unui ingredient nou direct din editorul de retete, atunci cand acesta nu exista.")
    h.bullet(doc, "Functia de resetare a parolei din ecranul de autentificare.")
    h.bullet(doc, "Atasarea de fotografii la retete si salvarea ciornelor (drafts).")
    h.bullet(doc, "Conversia aplicatiei intr-un executabil (.exe) distribuibil si un ecran de incarcare.")
    h.bullet(doc, "Un strat de conversie intre unitati de masura pentru camara si lista de cumparaturi.")

    h.heading(doc, "Comentariu personal", level=2)
    h.body(doc,
        "Raportat la obiectivele enuntate in introducere, consider ca stagiul a fost atins in proportie "
        "mare: produsul este coerent, sigur si demonstrabil, iar deciziile de proiectare sunt documentate si "
        "justificate. Am invatat ca o arhitectura buna se vede mai ales in lucrurile care nu se intampla — "
        "absenta injectiei SQL, absenta stergerilor accidentale in cascada, absenta blocarii interfetei la "
        "apelurile de retea. De asemenea, folosirea disciplinata a versionarii si a documentarii s-a dovedit "
        "esentiala pentru a duce un proiect peste mai multe sesiuni de lucru, fara a pierde contextul. "
        "Elementele ramase nefinalizate nu sunt esecuri, ci puncte de plecare clare pentru o versiune urmatoare.")


# =========================================================================== #
#  6. BIBLIOGRAFIE (WEBOGRAFIE)
# =========================================================================== #

def bibliografie(doc, h):
    h.heading(doc, "Bibliografie (Webografie)", level=1)
    surse = [
        "Microsoft. Documentatia SQL Server si limbajul Transact-SQL. https://learn.microsoft.com/en-us/sql/",
        "Microsoft. OPENJSON (Transact-SQL). https://learn.microsoft.com/en-us/sql/t-sql/functions/openjson-transact-sql",
        "Microsoft. .NET si WPF (Windows Presentation Foundation). https://learn.microsoft.com/en-us/dotnet/desktop/wpf/",
        "Microsoft. CommunityToolkit.Mvvm (MVVM Toolkit). https://learn.microsoft.com/en-us/dotnet/communitytoolkit/mvvm/",
        "Microsoft. Microsoft.Data.SqlClient. https://learn.microsoft.com/en-us/sql/connect/ado-net/",
        "Microsoft. Dependency injection in .NET. https://learn.microsoft.com/en-us/dotnet/core/extensions/dependency-injection",
        "Dapper — a simple object mapper for .NET. https://github.com/DapperLib/Dapper",
        "BCrypt.Net-Next — hashing de parole pentru .NET. https://github.com/BcryptNet/bcrypt.net",
        "ClosedXML — citire/scriere fisiere Excel (.xlsx). https://github.com/ClosedXML/ClosedXML",
        "Docker. Documentatia oficiala. https://docs.docker.com/",
        "Microsoft. Imaginea de container SQL Server. https://hub.docker.com/_/microsoft-mssql-server",
        "OWASP. SQL Injection Prevention Cheat Sheet. https://cheatsheetseries.owasp.org/",
        "OWASP. Password Storage Cheat Sheet. https://cheatsheetseries.owasp.org/",
        "Obsidian — aplicatie de gestiune a cunostintelor. https://obsidian.md/",
        "Git. Documentatia oficiala (Pro Git). https://git-scm.com/doc",
        "Anthropic. Claude Code — documentatie. https://docs.claude.com/en/docs/claude-code",
    ]
    for s in surse:
        h.numbered(doc, s)
