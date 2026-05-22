"""
Ajutoare de formatare pentru raportul de practica.

Toate regulile stricte de formatare cerute de indrumar sunt centralizate aici:
  - Pagina A4, margini: stanga 30 mm, sus 20 mm, jos 20 mm, dreapta 10 mm.
  - Corp text: Times New Roman 12, justify, spatiere 1.5 randuri.
  - Titluri: Times New Roman 14, aldin (bold), centrat.
  - Cod sursa: Courier New 10, aliniere stanga, spatiere 1 rand.
  - Legenda figurilor: sub imagine, centrat ("Figura N - ...").
  - Numerotarea paginilor: jos, centrat.
  - Cuprins: camp TOC nativ Word (se actualizeaza cu F9 / la deschidere).
"""

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt, RGBColor

BODY_FONT = "Times New Roman"
CODE_FONT = "Courier New"
BLACK = RGBColor(0, 0, 0)

# Numarator global pentru figuri (legendele cer "Figura N").
_FIGURE_COUNTER = {"n": 0}


# --------------------------------------------------------------------------- #
#  Stiluri si setari de document
# --------------------------------------------------------------------------- #

def _set_font(font, name, size_pt, bold=False, italic=False, color=None):
    font.name = name
    # rFonts pentru a forta fontul pe toate seturile de caractere (inclusiv latin).
    rpr = font.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color is not None:
        font.color.rgb = color


def configure_styles(doc):
    """Aplica toate regulile de formatare pe stilurile documentului."""
    styles = doc.styles

    # Corp text: TNR 12, justify, 1.5 randuri.
    normal = styles["Normal"]
    _set_font(normal.font, BODY_FONT, 12, color=BLACK)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)

    # Titluri (Heading 1..3): TNR 14, aldin, centrat, negru.
    # Raman stiluri de tip "Heading" ca sa fie prinse de campul TOC.
    for level in (1, 2, 3):
        h = styles[f"Heading {level}"]
        _set_font(h.font, BODY_FONT, 14, bold=True, color=BLACK)
        hpf = h.paragraph_format
        hpf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hpf.space_before = Pt(12)
        hpf.space_after = Pt(6)
        hpf.keep_with_next = True

    # Stil cod sursa: Courier New 10, stanga, 1 rand.
    if "Cod" not in [s.name for s in styles]:
        from docx.enum.style import WD_STYLE_TYPE
        code = styles.add_style("Cod", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Cod"]
    _set_font(code.font, CODE_FONT, 10, color=BLACK)
    cpf = code.paragraph_format
    cpf.line_spacing = 1.0
    cpf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cpf.space_after = Pt(0)
    cpf.space_before = Pt(0)

    # Stil legenda figura/tabel: centrat.
    if "Legenda" not in [s.name for s in styles]:
        from docx.enum.style import WD_STYLE_TYPE
        cap = styles.add_style("Legenda", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Legenda"]
    _set_font(cap.font, BODY_FONT, 11, italic=True, color=BLACK)
    capf = cap.paragraph_format
    capf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    capf.space_before = Pt(2)
    capf.space_after = Pt(10)
    capf.line_spacing = 1.0


def set_a4_margins(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.right_margin = Mm(10)


def enable_auto_update_fields(doc):
    """Cere Word/LibreOffice sa actualizeze campurile (TOC, PAGE) la deschidere."""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)


# --------------------------------------------------------------------------- #
#  Numerotarea paginilor (subsol, centrat)
# --------------------------------------------------------------------------- #

def _add_field(paragraph, instr):
    """Insereaza un camp Word (ex. PAGE) intr-un paragraf."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_el)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_page_number_footer(section, show=True):
    """Subsol cu numarul paginii, centrat. show=False => subsol gol (foaia de titlu)."""
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if show:
        _add_field(p, " PAGE ")
        for r in p.runs:
            _set_font(r.font, BODY_FONT, 11, color=BLACK)


# --------------------------------------------------------------------------- #
#  Continut: titluri, paragrafe, cod, figuri, tabele, cuprins
# --------------------------------------------------------------------------- #

def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    return p


def code_block(doc, text, max_lines=None):
    """Bloc de cod: o linie = un paragraf in stilul 'Cod' (pagineaza corect)."""
    lines = text.rstrip("\n").split("\n")
    truncated = False
    if max_lines is not None and len(lines) > max_lines:
        lines = lines[:max_lines]
        truncated = True
    for line in lines:
        p = doc.add_paragraph(style="Cod")
        # spatiile la inceput de linie se pastreaza (indentarea codului)
        run = p.add_run(line if line else "")
        run.font.name = CODE_FONT
    if truncated:
        p = doc.add_paragraph(style="Cod")
        p.add_run("    [... fragment trunchiat; listing complet in Anexa A1 ...]")


def figure_placeholder(doc, caption, height_hint="120 mm"):
    """
    Chenar gol pentru o captura de ecran (de adaugat de pe VM-ul Windows),
    urmat de legenda centrata "Figura N - ...".
    """
    _FIGURE_COUNTER["n"] += 1
    n = _FIGURE_COUNTER["n"]
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = box.add_run(f"[ Spatiu pentru imagine — inseram aici captura ({height_hint}) ]")
    _set_font(r.font, BODY_FONT, 11, italic=True, color=RGBColor(0x80, 0x80, 0x80))
    _box_border(box)
    cap = doc.add_paragraph(style="Legenda")
    cap.add_run(f"Figura {n} — {caption}")
    return n


def _box_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "8")
        e.set(qn("w:color"), "AAAAAA")
        pbdr.append(e)
    pPr.append(pbdr)


def simple_table(doc, headers, rows, caption=None):
    """Tabel cu antet + randuri. Optional o legenda dedesubt (centrata)."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        _set_font(run.font, BODY_FONT, 11, bold=True, color=BLACK)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            _set_font(run.font, BODY_FONT, 11, color=BLACK)
    if caption:
        cap = doc.add_paragraph(style="Legenda")
        cap.add_run(caption)
    return table


def toc_field(doc):
    """Camp TOC nativ: niveluri 1-3, cu hyperlink-uri, actualizabil (F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:r")
    ptext = OxmlElement("w:t")
    ptext.text = "Cuprinsul se genereaza automat — clic dreapta › Update Field (sau F9) in Word."
    placeholder.append(ptext)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)
    return p


def page_break(doc):
    doc.add_page_break()


def new_section_page(doc):
    """Sectiune noua incepand pe pagina noua (pentru subsoluri/numerotare distincte)."""
    return doc.add_section(WD_SECTION.NEW_PAGE)
